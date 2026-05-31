from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# добавление знака 
def add_watermark_docx(
    input_path: str,
    output_path: str,
    watermark_text: str,
    doc_uuid: str
):
    try:
        doc = Document(input_path)

        for section in doc.sections:
            header = section.header
            if header.paragraphs:
                paragraph = header.paragraphs[0]
            else:
                paragraph = header.add_paragraph()

            paragraph.clear()
            run = paragraph.add_run(watermark_text)
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(180, 180, 180)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            footer = section.footer
            if footer.paragraphs:
                footer_paragraph = footer.paragraphs[0]
            else:
                footer_paragraph = footer.add_paragraph()

            footer_paragraph.clear()
            footer_run = footer_paragraph.add_run(watermark_text)
            footer_run.font.size = Pt(10)
            footer_run.bold = False
            footer_run.font.color.rgb = RGBColor(200, 200, 200)
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        core = doc.core_properties
        core.comments = f"DOCUMENT_ID:{doc_uuid}"
        core.category = "watermark"
        core.subject = "protected_by_bot"

        doc.save(output_path)
        return True
    except Exception as e:
        print(f'Ошибка при обработке DOCX: {e}')
        return False

# проверка наличия знака в docx 
def check_watermark_docx(file_path: str, expected_uuid: str) -> dict:
    try:
        doc = Document(file_path)

        core = doc.core_properties
        meta_match = False
        
        if core.comments and expected_uuid in core.comments:
            meta_match = True

        visible_match = False
        for section in doc.sections:
            if section.header.paragraphs:
                header_text = section.header.paragraphs[0].text.strip()
                if header_text:
                    visible_match = True
                    break

        return {
            "visible_watermark": visible_match,
            "hidden_watermark": meta_match,
            "watermark_found": visible_match or meta_match
        }
    
    except Exception as e:
        print(f'Ошибка при проверке DOCX: {e}')
        return {
            "visible_watermark": False,
            "hidden_watermark": False,
            "watermark_found": False
        }